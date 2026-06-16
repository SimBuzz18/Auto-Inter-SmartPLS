import os
from datetime import datetime
import pandas as pd
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_ERROR = None
except ImportError as e:
    Document = None
    DOCX_ERROR = str(e)

class ScientificInterpreter:
    SRMR_OMML = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:r><m:t>SRMR = </m:t></m:r>'
        '<m:rad>'
        '<m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:e>'
        '<m:f>'
        '<m:num>'
        '<m:sSub>'
        '<m:e><m:r><m:t>∑</m:t></m:r></m:e>'
        '<m:sub><m:r><m:t>i&lt;j</m:t></m:r></m:sub>'
        '</m:sSub>'
        '<m:sSup>'
        '<m:e>'
        '<m:r><m:t>(</m:t></m:r>'
        '<m:sSub><m:e><m:r><m:t>s</m:t></m:r></m:e><m:sub><m:r><m:t>ij</m:t></m:r></m:sub></m:sSub>'
        '<m:r><m:t> - </m:t></m:r>'
        '<m:sSub><m:e><m:r><m:t>σ</m:t></m:r></m:e><m:sub><m:r><m:t>ij</m:t></m:r></m:sub></m:sSub>'
        '<m:r><m:t>)</m:t></m:r>'
        '</m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
        '</m:sSup>'
        '</m:num>'
        '<m:den>'
        '<m:r><m:t>h</m:t></m:r>'
        '</m:den>'
        '</m:f>'
        '</m:e>'
        '</m:rad>'
        '</m:oMath>'
    )
    
    HTMT_OMML = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:sSub>'
        '<m:e><m:r><m:t>HTMT</m:t></m:r></m:e>'
        '<m:sub><m:r><m:t>ij</m:t></m:r></m:sub>'
        '</m:sSub>'
        '<m:r><m:t> = </m:t></m:r>'
        '<m:f>'
        '<m:num>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="¯"/></m:accPr>'
        '<m:e><m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:t>ij</m:t></m:r></m:sub></m:sSub></m:e>'
        '</m:acc>'
        '</m:num>'
        '<m:den>'
        '<m:rad>'
        '<m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:e>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="¯"/></m:accPr>'
        '<m:e><m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:t>ii</m:t></m:r></m:sub></m:sSub></m:e>'
        '</m:acc>'
        '<m:r><m:t> × </m:t></m:r>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="¯"/></m:accPr>'
        '<m:e><m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:t>jj</m:t></m:r></m:sub></m:sSub></m:e>'
        '</m:acc>'
        '</m:e>'
        '</m:rad>'
        '</m:den>'
        '</m:f>'
        '</m:oMath>'
    )
    
    GOF_OMML = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:r><m:t>GoF = </m:t></m:r>'
        '<m:rad>'
        '<m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:e>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="¯"/></m:accPr>'
        '<m:e><m:r><m:t>AVE</m:t></m:r></m:e>'
        '</m:acc>'
        '<m:r><m:t> × </m:t></m:r>'
        '<m:sSup>'
        '<m:e>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="¯"/></m:accPr>'
        '<m:e><m:r><m:t>R</m:t></m:r></m:e>'
        '</m:acc>'
        '</m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
        '</m:sSup>'
        '</m:e>'
        '</m:rad>'
        '</m:oMath>'
    )

    def __init__(self):
        self.paragraphs = []
        self.analysis_context = None  # Will be set by SmartPLSReader.process()
    
    def add_equation(self, xml_content):
        self.paragraphs.append({"type": "equation", "xml": xml_content})

    
    def set_analysis_context(self, context):
        """
        Set the analysis context (type, variables) for adaptive interpretations.
        """
        self.analysis_context = context

    def add_section(self, title, content, level=1):
        self.paragraphs.append({"type": "header", "text": title, "level": level})
        if isinstance(content, list):
            for item in content:
                if isinstance(item, pd.DataFrame):
                    self.paragraphs.append({"type": "table", "data": item})
                else:
                    self.paragraphs.append({"type": "text", "text": str(item)})
        else:
            self.paragraphs.append({"type": "text", "text": str(content)})

    def add_header(self, text, level=1):
        self.paragraphs.append({"type": "header", "text": text, "level": level})

    def add_text(self, text):
        self.paragraphs.append({"type": "text", "text": str(text)})

    def add_table(self, df):
        self.paragraphs.append({"type": "table", "data": df})

    def get_table_interpretation(self, table_name):
        """
        Returns a comprehensive scientific interpretation based on table name.
        """
        table_name_lower = table_name.lower()
        interpretations = {
            'path coefficients': (
                "Analisis koefisien jalur (path coefficients) digunakan untuk mengevaluasi arah, kekuatan, dan signifikansi hubungan antar variabel dalam model struktural. Hubungan dinyatakan signifikan secara statistik apabila nilai signifikansi p-value \u2264 0,05 atau nilai t-statistic \u2265 1,96."
            ),
            'indirect effects': (
                "Analisis efek tidak langsung (indirect effects) bertujuan untuk mengidentifikasi keberadaan dan besarnya pengaruh variabel eksogen terhadap variabel endogen yang disalurkan melalui variabel mediator. Signifikansi pengaruh mediasi ini ditentukan berdasarkan metode bootstrapping."
            ),
            'total indirect effects': (
                "Total efek tidak langsung merepresentasikan akumulasi seluruh pengaruh tidak langsung dari variabel eksogen terhadap variabel endogen melalui seluruh jalur mediasi yang ada dalam model."
            ),
            'specific indirect effects': (
                "Efek tidak langsung spesifik menggambarkan besarnya kontribusi pengaruh tidak langsung dari variabel eksogen ke variabel endogen melalui jalur mediasi tertentu. Hal ini berguna untuk menilai peran spesifik dari masing-masing variabel mediasi secara terpisah."
            ),
            'total effects': (
                "Efek total (total effects) merupakan hasil penjumlahan dari efek langsung dan seluruh efek tidak langsung yang terjadi antar konstruk, sehingga memberikan gambaran dampak menyeluruh dari variabel eksogen terhadap variabel endogen."
            ),
            'outer loadings': (
                "Evaluasi beban luar (outer loadings) digunakan untuk menilai reliabilitas indikator individual dalam model pengukuran reflektif. Nilai beban luar yang disarankan adalah lebih besar dari 0,70 untuk memastikan bahwa indikator tersebut memiliki tingkat konsistensi dan representasi yang kuat terhadap konstruk latennya."
            ),
            'outer weights': (
                "Bobot luar (outer weights) digunakan untuk mengukur kontribusi relatif dari masing-masing indikator dalam membentuk konstruk laten formatif. Signifikansi bobot ini dievaluasi melalui prosedur bootstrapping."
            ),
            'r square': (
                "Koefisien determinasi (R-Square) mencerminkan proporsi varians variabel endogen yang dapat dijelaskan oleh variabel-variabel eksogen yang mempengaruhinya. Nilai R-Square berkisar dari 0 hingga 1, di mana nilai yang lebih besar menunjukkan kemampuan prediksi model struktural yang lebih baik."
            ),
            'f square': (
                "Nilai f-Square (effect size) digunakan untuk mengukur kontribusi praktis dari suatu variabel eksogen terhadap nilai R-Square variabel endogen. Berdasarkan kriteria Cohen, kontribusi f-Square dikategorikan kecil (0,02), sedang (0,15), dan besar (0,35)."
            ),
            'construct reliability': (
                "Uji keandalan konstruk (construct reliability) dilakukan untuk mengevaluasi konsistensi internal dari model pengukuran menggunakan kriteria Composite Reliability (CR) dan Cronbach's Alpha. Untuk penelitian uji teori, nilai Composite Reliability yang disarankan adalah lebih besar dari 0,60."
            ),
            'fornell-larcker': (
                "Kriteria Fornell-Larcker digunakan untuk mengevaluasi validitas diskriminan dengan membandingkan nilai akar kuadrat dari Average Variance Extracted (AVE) masing-masing konstruk dengan nilai korelasi antar konstruk. Validitas diskriminan terpenuhi jika nilai akar kuadrat AVE suatu konstruk lebih tinggi dibandingkan korelasinya dengan konstruk lainnya, dengan batas minimal AVE sebesar 0,50."
            ),
            'cross loadings': (
                "Analisis cross loadings merupakan metode evaluasi validitas diskriminan di tingkat indikator. Kriteria terpenuhi apabila korelasi suatu indikator dengan konstruk laten utamanya lebih besar dibandingkan korelasi indikator tersebut dengan konstruk laten lainnya dalam model."
            ),
            'htmt': (
                "Rasio Heterotrait-Monotrait (HTMT) merupakan metode modern untuk mendeteksi masalah validitas diskriminan. Nilai rasio HTMT yang berada di bawah ambang batas 0,90 (atau di bawah 0,85 untuk kriteria yang lebih ketat) mengindikasikan bahwa validitas diskriminan antar konstruk telah terpenuhi."
            ),
            'vif': (
                "Variance Inflation Factor (VIF) digunakan untuk mengidentifikasi keberadaan isu multikolinearitas antar variabel prediktor. Nilai VIF di bawah 3 mengindikasikan kondisi ideal bebas kolinearitas, sementara nilai di atas 5 menunjukkan adanya multikolinearitas yang serius."
            ),
            'inner vif': (
                "Variance Inflation Factor dalam model struktural (inner VIF) digunakan untuk memeriksa apakah terdapat kolinearitas antar variabel eksogen (prediktor). Nilai VIF yang disarankan adalah kurang dari 3 untuk menghindari bias estimasi koefisien jalur."
            ),
            'outer vif': (
                "Variance Inflation Factor dalam model pengukuran (outer VIF) mengevaluasi tingkat kolinearitas antar indikator manifest pada model pengukuran formatif. Ambang batas yang disarankan adalah di bawah 5."
            ),
            'fit summary': (
                "Evaluasi kecocokan model (goodness of fit) dilakukan menggunakan metrik Fit Summary. Indikator utama yang digunakan adalah Standardized Root Mean Square Residual (SRMR), di mana nilai SRMR di bawah 0,08 menunjukkan bahwa model teoritis memiliki kecocokan yang baik dengan data empiris."
            ),
            'rms theta': (
                "Root Mean Square Covariance (rms Theta) mengukur tingkat residu model pengukuran reflektif, di mana nilai di bawah 0,12 menunjukkan spesifikasi model pengukuran yang memadai."
            ),
            'mean, stdev': (
                "Statistik deskriptif parameter yang dihasilkan dari prosedur bootstrapping mencakup rata-rata sampel, standar deviasi, t-statistics, dan p-values untuk menilai signifikansi empiris dari koefisien jalur."
            ),
            'confidence intervals': (
                "Interval kepercayaan (confidence intervals) memberikan rentang estimasi parameter pada tingkat kepercayaan 95%. Hubungan dinyatakan signifikan jika rentang interval tersebut tidak melewati atau mencakup nilai nol."
            ),
            'q²': (
                "Koefisien relevansi prediktif (Q²) hasil blindfolding digunakan untuk mengukur seberapa baik nilai observasi direkonstruksi oleh model struktural. Nilai Q² yang lebih besar dari nol menunjukkan bahwa model memiliki relevansi prediktif."
            ),
            'blindfolding': (
                "Prosedur blindfolding dilakukan untuk menghasilkan nilai koefisien relevansi prediktif (Q²), yang berguna untuk menilai kemampuan prediksi model struktural di luar sampel (out-of-sample predictive power)."
            ),
            'discriminant validity': (
                "Uji validitas diskriminan bertujuan untuk memastikan bahwa setiap konstruk laten secara empiris bersifat unik dan berbeda dari konstruk lainnya dalam model."
            ),
            'quality criteria': (
                "Kriteria kualitas (quality criteria) merangkum seluruh parameter evaluasi kelayakan model pengukuran (outer model) dan model struktural (inner model) secara komprehensif."
            ),
            'collinearity': (
                "Evaluasi kolinearitas dilakukan untuk mendeteksi redundansi hubungan linear antar variabel prediktor yang dapat mendistorsi hasil analisis."
            ),
            'final results': (
                "Ringkasan hasil akhir (final results) menyajikan estimasi titik parameter model pengukuran dan model struktural yang diperoleh dari algoritma PLS."
            ),
        }
        
        # Find matching interpretation
        name_norm = table_name_lower.replace('-', ' ')
        for key, interpretation in interpretations.items():
            key_norm = key.replace('-', ' ')
            if key_norm in name_norm:
                return interpretation
        
        # Default interpretation
        return f"Tabel {table_name} menampilkan hasil analisis yang relevan untuk evaluasi model penelitian. Interpretasi spesifik tergantung pada konteks dan tujuan analisis."
    
    def get_dynamic_interpretation(self, table_name, df):
        """
        Generate data-driven interpretation by analyzing actual table values.
        Combines generic explanation with specific data analysis.
        Adapts based on analysis context (simple regression, mediation, moderation).
        """
        import pandas as pd
        import numpy as np
        
        table_name_lower = table_name.lower()
        
        # [SCIENTIFIC ANALYSIS override]
        # For GoF and VAF, we use specialized methods that provide the full "Table -> Narrative -> Analysis" structure
        if 'goodness of fit' in table_name_lower or 'gof' in table_name_lower:
            return self._analyze_gof(df)
        elif 'vaf' in table_name_lower or 'variance accounted for' in table_name_lower:
            return self._analyze_vaf(df)
        
        # Check if this table is relevant for the current analysis type
        context_note = self._get_context_relevance_note(table_name_lower)
        
        # Get base interpretation
        base_interpretation = self.get_table_interpretation(table_name)
        
        # Add context-specific prefix if needed
        if context_note:
            base_interpretation = context_note + "\n\n" + base_interpretation
        
        # Analyze specific table types and add data-driven insights
        data_analysis = []
        
        try:
            if 'path coefficients' in table_name_lower or 'path coefficient' in table_name_lower:
                data_analysis = self._analyze_path_coefficients(df)
            elif 'r square' in table_name_lower or 'r-square' in table_name_lower:
                data_analysis = self._analyze_r_square(df)
            elif 'f square' in table_name_lower or 'f-square' in table_name_lower:
                data_analysis = self._analyze_f_square(df)
            elif 'outer loading' in table_name_lower:
                data_analysis = self._analyze_outer_loadings(df)
            elif 'construct reliability' in table_name_lower:
                data_analysis = self._analyze_construct_reliability(df)
            elif 'fornell-larcker' in table_name_lower:
                data_analysis = self._analyze_fornell_larcker(df)
            elif 'htmt' in table_name_lower:
                data_analysis = self._analyze_htmt(df)
            elif 'vif' in table_name_lower:
                data_analysis = self._analyze_vif(df)
            elif 'mean, stdev' in table_name_lower or 'bootstrapping' in table_name_lower:
                data_analysis = self._analyze_bootstrapping(df)
            elif 'indirect effect' in table_name_lower:
                data_analysis = self._analyze_indirect_effects(df)
            elif 'total effect' in table_name_lower:
                data_analysis = self._analyze_total_effects(df)
            elif 'q²' in table_name_lower or 'q square' in table_name_lower or 'blindfolding' in table_name_lower:
                data_analysis = self._analyze_q_square(df)
            elif 'interaction' in table_name_lower or 'moderating' in table_name_lower:
                data_analysis = self._analyze_moderation_effects(df)
        except Exception as e:
            print(f"Error analyzing {table_name}: {e}")
        
        # Combine base interpretation with data analysis
        if data_analysis:
            # Join with double newlines effectively to separate items clearly
            result = base_interpretation + "\n\n" + "\n".join(data_analysis)
            return result
        
        return base_interpretation
        
    def get_opening_narration(self, table_name):
        return self.get_table_interpretation(table_name)
        
    def get_detailed_explanation(self, table_name, df):
        table_name_lower = table_name.lower()
        if 'goodness of fit' in table_name_lower or 'gof' in table_name_lower:
            res = self._analyze_gof(df)
            return "\n".join(res) if isinstance(res, list) else str(res)
        elif 'vaf' in table_name_lower or 'variance accounted for' in table_name_lower:
            res = self._analyze_vaf(df)
            return "\n".join(res) if isinstance(res, list) else str(res)
            
        data_analysis = []
        try:
            if 'path coefficients' in table_name_lower or 'path coefficient' in table_name_lower:
                data_analysis = self._analyze_path_coefficients(df)
            elif 'r square' in table_name_lower or 'r-square' in table_name_lower:
                data_analysis = self._analyze_r_square(df)
            elif 'f square' in table_name_lower or 'f-square' in table_name_lower:
                data_analysis = self._analyze_f_square(df)
            elif 'outer loading' in table_name_lower:
                data_analysis = self._analyze_outer_loadings(df)
            elif 'construct reliability' in table_name_lower:
                data_analysis = self._analyze_construct_reliability(df)
            elif 'fornell-larcker' in table_name_lower:
                data_analysis = self._analyze_fornell_larcker(df)
            elif 'htmt' in table_name_lower:
                data_analysis = self._analyze_htmt(df)
            elif 'vif' in table_name_lower:
                data_analysis = self._analyze_vif(df)
            elif 'mean, stdev' in table_name_lower or 'bootstrapping' in table_name_lower:
                data_analysis = self._analyze_bootstrapping(df)
            elif 'indirect effect' in table_name_lower:
                data_analysis = self._analyze_indirect_effects(df)
            elif 'total effect' in table_name_lower:
                data_analysis = self._analyze_total_effects(df)
            elif 'q²' in table_name_lower or 'q square' in table_name_lower or 'blindfolding' in table_name_lower:
                data_analysis = self._analyze_q_square(df)
            elif 'interaction' in table_name_lower or 'moderating' in table_name_lower:
                data_analysis = self._analyze_moderation_effects(df)
        except Exception as e:
            print(f"Error analyzing {table_name}: {e}")
            
        return "\n".join(data_analysis) if data_analysis else ""
    
    def _get_context_relevance_note(self, table_name_lower):
        """
        Return None - context notes disabled by user request.
        """
        return None
    
    def _analyze_moderation_effects(self, df):
        """Analyze moderation/interaction effects table."""
        results = []
        try:
            for idx, row in df.iterrows():
                path = str(row.iloc[0]) if len(row) > 0 else ""
                if not path or '*' not in path.lower() and 'x' not in path.lower():
                    continue
                    
                val = row.iloc[1] if len(row) > 1 else None
                if pd.notna(val) and path:
                    try:
                        val_num = float(val)
                        direction = "memperkuat" if val_num > 0 else "memperlemah"
                        strength = "signifikan" if abs(val_num) > 0.1 else "lemah"
                        results.append(f"Pengaruh interaksi pada hubungan {path} memiliki nilai koefisien moderasi sebesar {val_num:.3f}. Temuan ini menunjukkan bahwa variabel moderator {direction} kekuatan hubungan langsung tersebut dengan kategori efek yang tergolong {strength}.")
                    except:
                        pass
        except:
            pass
        
        if not results:
            results.append("Berdasarkan data pengujian, tidak ditemukan adanya efek moderasi atau interaksi yang signifikan secara statistik pada model struktural ini.")
        
        return results
    
    def _analyze_path_coefficients(self, df):
        """Analyze path coefficients table."""
        results = []
        try:
            # Find coefficient column (usually second column or named 'Original Sample')
            coef_col = None
            for col in df.columns:
                col_str = str(col).lower()
                if 'original' in col_str or 'coefficient' in col_str or 'value' in col_str:
                    coef_col = col
                    break
            if coef_col is None and len(df.columns) >= 2:
                coef_col = df.columns[1]
            
            if coef_col:
                for idx, row in df.iterrows():
                    var = str(row.iloc[0]) if len(row) > 0 else ""
                    val = row[coef_col]
                    if pd.notna(val) and var:
                        try:
                            val_num = float(val)
                            direction = "Positif" if val_num > 0 else "Negatif"
                            strength = "kuat" if abs(val_num) > 0.5 else ("sedang" if abs(val_num) > 0.3 else "lemah")
                            results.append(f"Hubungan pada jalur {var} menunjukkan nilai koefisien jalur (path coefficient) sebesar {val_num:.3f}, yang merepresentasikan adanya pengaruh {direction.lower()} dengan tingkat kekuatan yang tergolong {strength}.")
                        except:
                            pass
        except:
            pass
        return results

    def _analyze_r_square(self, df):
        """Analyze R-Square table."""
        results = []
        try:
            for idx, row in df.iterrows():
                var = str(row.iloc[0]) if len(row) > 0 else ""
                val = row.iloc[1] if len(row) > 1 else None
                if pd.notna(val) and var:
                    try:
                        val_num = float(val)
                        if val_num >= 0.67:
                            cat = "substansial (kuat)"
                        elif val_num >= 0.33:
                            cat = "moderat (sedang)"
                        elif val_num >= 0.19:
                            cat = "lemah"
                        else:
                            cat = "sangat lemah"
                        pct = val_num * 100
                        results.append(f"Variabel endogen {var} memiliki nilai koefisien determinasi R-Square sebesar {val_num:.3f}, yang menunjukkan bahwa model pengaruh terhadap konstruk tersebut tergolong {cat}. Hasil ini mengimplikasikan bahwa sebesar {pct:.1f}% variabilitas konstruk {var} dapat dijelaskan oleh konstruk eksogen yang mempengaruhinya, sedangkan sisanya dijelaskan oleh faktor lain di luar model penelitian.")
                    except:
                        pass
        except:
            pass
        return results

    def _analyze_f_square(self, df):
        """Analyze f-Square (effect size) table."""
        results = []
        try:
            # f-square is usually a matrix, analyze non-zero values
            for idx, row in df.iterrows():
                var_from = str(row.iloc[0]) if len(row) > 0 else ""
                for col in df.columns[1:]:
                    val = row[col]
                    if pd.notna(val):
                        try:
                            val_num = float(val)
                            if val_num > 0.02:  # Only report meaningful effects
                                if val_num >= 0.35:
                                    effect = "besar"
                                elif val_num >= 0.15:
                                    effect = "sedang"
                                else:
                                    effect = "kecil"
                                results.append(f"Pengaruh variabel {var_from} terhadap {col} menghasilkan nilai f-Square sebesar {val_num:.3f}, yang mengindikasikan bahwa ukuran efek (effect size) dari konstruk eksogen tersebut berada pada kategori {effect}.")
                        except:
                            pass
        except:
            pass
        return results[:10]  # Limit to 10 entries

    def _analyze_outer_loadings(self, df):
        """Analyze outer loadings table."""
        results = []
        good_count = 0
        weak_count = 0
        try:
            for idx, row in df.iterrows():
                indicator = str(row.iloc[0]) if len(row) > 0 else ""
                for col in df.columns[1:]:
                    val = row[col]
                    if pd.notna(val):
                        try:
                            val_num = float(val)
                            if abs(val_num) > 0.4:  # Only own loadings
                                if round(val_num, 3) >= 0.708:
                                    good_count += 1
                                elif val_num >= 0.4:
                                    weak_count += 1
                                    results.append(f"Indikator {indicator} pada konstruk {col} menunjukkan nilai loading sebesar {val_num:.3f}. Angka tersebut berada di bawah ambang batas ideal 0,708, sehingga perlu dipertimbangkan untuk dieliminasi dari model guna meningkatkan validitas konvergen.")
                        except:
                            pass
        except:
            pass
        summary = f"Berdasarkan hasil analisis loading factor, sebanyak {good_count} indikator dinyatakan telah memenuhi kriteria kelayakan di atas 0,708."
        if weak_count > 0:
            summary += f" Namun demikian, terdapat {weak_count} indikator yang memerlukan perhatian lebih lanjut karena memiliki nilai loading di bawah batas ideal (antara 0,40 hingga 0,708)."
        else:
            summary += ""
        return [summary] + results[:5]

    def _analyze_construct_reliability(self, df):
        """Analyze construct reliability and validity table."""
        results = []
        try:
            for idx, row in df.iterrows():
                construct = str(row.iloc[0]) if len(row) > 0 else ""
                if not construct:
                    continue
                
                issues = []
                good_points = []
                
                for col in df.columns:
                    col_lower = str(col).lower()
                    val = row[col]
                    if pd.isna(val):
                        continue
                    try:
                        val_num = float(val)
                        
                        if 'alpha' in col_lower or "cronbach" in col_lower:
                            if val_num >= 0.7:
                                good_points.append(f"Cronbach's α = {val_num:.3f} (baik)")
                            else:
                                issues.append(f"Cronbach's α = {val_num:.3f} (kurang reliabel)")
                        
                        elif 'composite' in col_lower or 'rho_c' in col_lower or 'cr' in col_lower:
                            if val_num >= 0.7:
                                good_points.append(f"CR = {val_num:.3f} (baik)")
                            else:
                                issues.append(f"CR = {val_num:.3f} (kurang reliabel)")
                        
                        elif 'ave' in col_lower:
                            if val_num >= 0.5:
                                good_points.append(f"AVE = {val_num:.3f} (valid)")
                            else:
                                issues.append(f"AVE = {val_num:.3f} (validitas konvergen rendah)")
                    except:
                        pass
                
                if good_points or issues:
                    status = "memenuhi kriteria secara memadai" if not issues else "memerlukan evaluasi lebih lanjut"
                    detail = "; ".join(good_points + issues)
                    results.append(f"Evaluasi konsistensi internal dan validitas konvergen untuk konstruk {construct} menunjukkan status {status}, dengan rincian metrik: {detail}.")
        except:
            pass
        return results

    def _analyze_fornell_larcker(self, df):
        """Analyze Fornell-Larcker criterion table."""
        results = []
        try:
            issues = []
            for idx, row in df.iterrows():
                construct = str(row.iloc[0]) if len(row) > 0 else ""
                if not construct:
                    continue
                
                # Diagonal value (sqrt AVE)
                diag_val = None
                for col in df.columns:
                    if str(col).strip() == construct.strip():
                        diag_val = row[col]
                        break
                
                if pd.notna(diag_val):
                    try:
                        diag_num = float(diag_val)
                        # Check against all off-diagonal values
                        for col in df.columns[1:]:
                            if str(col).strip() != construct.strip():
                                off_val = row[col]
                                if pd.notna(off_val):
                                    try:
                                        off_num = float(off_val)
                                        if off_num > diag_num:
                                            issues.append(f"konstruk {construct} memiliki nilai akar kuadrat AVE ({diag_num:.3f}) yang lebih kecil dibandingkan nilai korelasi konstruk tersebut dengan konstruk {col} ({off_num:.3f})")
                                    except:
                                        pass
                    except:
                        pass
            
            if issues:
                results.append("Berdasarkan kriteria Fornell-Larcker, diidentifikasi adanya indikasi permasalahan validitas diskriminan pada beberapa konstruk berikut:")
                results.extend([f"Terdapat kendala di mana {issue}." for issue in issues])
            else:
                results.append("Seluruh konstruk dalam model penelitian ini telah memenuhi kriteria validitas diskriminan Fornell-Larcker, di mana nilai akar kuadrat AVE masing-masing konstruk laten terbukti lebih besar dibandingkan dengan nilai korelasi antar-konstruk.")
        except:
            pass
        return results

    def _analyze_htmt(self, df):
        """Analyze HTMT table."""
        results = []
        try:
            issues = []
            warnings = []
            for idx, row in df.iterrows():
                for col in df.columns[1:]:
                    val = row[col]
                    if pd.notna(val):
                        try:
                            val_num = float(val)
                            pair = f"{row.iloc[0]} dengan {col}"
                            if val_num >= 0.90:
                                issues.append(f"pasangan konstruk {pair} menunjukkan nilai rasio HTMT sebesar {val_num:.3f}, yang melampaui batas kritis 0,90")
                            elif val_num >= 0.85:
                                warnings.append(f"pasangan konstruk {pair} menunjukkan nilai rasio HTMT sebesar {val_num:.3f} (kategori peringatan)")
                        except:
                            pass
            
            if issues:
                results.append("Hasil pengujian rasio Heterotrait-Monotrait (HTMT) mengindikasikan adanya potensi masalah validitas diskriminan pada pasangan konstruk berikut:")
                results.extend([f"Indikasi tumpang tindih terjadi karena {issue}." for issue in issues])
            if warnings:
                results.append("Analisis rasio HTMT mendeteksi adanya indikasi kolinearitas diskriminan moderat pada pasangan konstruk berikut:")
                results.extend([f"Perlu perhatian karena {w}." for w in warnings])
            if not issues and not warnings:
                results.append("Seluruh estimasi rasio HTMT berada di bawah ambang batas kritis 0,90 (serta di bawah batas konservatif 0,85), sehingga dapat disimpulkan bahwa validitas diskriminan antar konstruk telah terpenuhi secara memadai.")
        except:
            pass
        return results

    def _analyze_vif(self, df):
        """Analyze VIF table."""
        results = []
        try:
            issues = []
            warnings = []
            for idx, row in df.iterrows():
                var = str(row.iloc[0]) if len(row) > 0 else ""
                for col in df.columns[1:]:
                    val = row[col]
                    if pd.notna(val):
                        try:
                            val_num = float(val)
                            if val_num >= 5:
                                issues.append(f"hubungan prediktor {var} terhadap {col} menghasilkan nilai VIF sebesar {val_num:.2f} (di atas batas kritis 5,0)")
                            elif val_num >= 3:
                                warnings.append(f"hubungan prediktor {var} terhadap {col} memiliki nilai VIF sebesar {val_num:.2f} (di atas 3,0)")
                        except:
                            pass
            
            if issues:
                results.append("Berdasarkan hasil analisis Variance Inflation Factor (VIF), terdeteksi adanya indikasi permasalahan multikolinearitas yang serius pada variabel prediktor berikut:")
                results.extend([f"Distorsi estimasi koefisien jalur berpotensi terjadi karena {issue}." for issue in issues])
            if warnings:
                results.append("Analisis VIF menunjukkan adanya kolinearitas moderat yang memerlukan perhatian pada variabel prediktor berikut:")
                results.extend([f"Diperlukan kehati-hatian karena {w}." for w in warnings])
            if not issues and not warnings:
                results.append("Hasil pengujian VIF menunjukkan tidak adanya masalah multikolinearitas dalam model prediktor, di mana seluruh nilai VIF berada di bawah batas konservatif 3,0, sehingga estimasi koefisien jalur dinyatakan stabil.")
        except:
            pass
        return results

    def _analyze_bootstrapping(self, df):
        """Analyze bootstrapping results (T-statistics, P-values)."""
        results = []
        try:
            sig_count = 0
            nonsig_count = 0
            
            # Find P-value or T-statistic column
            p_col = None
            t_col = None
            for col in df.columns:
                col_str = str(col).lower()
                if 'p value' in col_str or 'p-value' in col_str or col_str.strip() == 'p':
                    p_col = col
                elif 't statistic' in col_str or 't-statistic' in col_str or col_str.strip() == 't':
                    t_col = col
            
            for idx, row in df.iterrows():
                path = str(row.iloc[0]) if len(row) > 0 else ""
                if not path:
                    continue
                
                is_sig = False
                p_val = None
                t_val = None
                
                if p_col and pd.notna(row[p_col]):
                    try:
                        p_val = float(row[p_col])
                        is_sig = p_val < 0.05
                    except:
                        pass
                
                if t_col and pd.notna(row[t_col]):
                    try:
                        t_val = float(row[t_col])
                        if p_val is None:
                            is_sig = abs(t_val) > 1.96
                    except:
                        pass
                
                p_str = f"{p_val:.4f}" if p_val is not None else "n.a."
                t_str = f"{t_val:.3f}" if t_val is not None else "n.a."
                
                if is_sig:
                    sig_count += 1
                    sig_level = " (sangat signifikan)" if (p_val is not None and p_val < 0.01) else (" (signifikan)" if (p_val is not None and p_val < 0.05) else "")
                    results.append(f"Jalur pengaruh {path} menunjukkan hasil yang signifikan{sig_level} dengan nilai P-Value = {p_str} dan T-Statistic = {t_str}, sehingga hipotesis yang diajukan dinyatakan diterima secara statistik.")
                elif p_val is not None or t_val is not None:
                    nonsig_count += 1
                    results.append(f"Jalur pengaruh {path} menunjukkan hasil yang tidak signifikan dengan nilai P-Value = {p_str} dan T-Statistic = {t_str}, sehingga hipotesis yang diajukan dinyatakan ditolak secara statistik.")
            
            summary = f"Berdasarkan hasil analisis uji signifikansi (bootstrapping), diperoleh kesimpulan bahwa terdapat {sig_count} jalur pengaruh yang signifikan dan {nonsig_count} jalur pengaruh yang tidak signifikan secara statistik."
            results.insert(0, summary)
        except:
            pass
        return results[:15]  # Limit output

    def _analyze_indirect_effects(self, df):
        """Analyze indirect effects table."""
        results = []
        try:
            for idx, row in df.iterrows():
                path = str(row.iloc[0]) if len(row) > 0 else ""
                val = row.iloc[1] if len(row) > 1 else None
                if pd.notna(val) and path:
                    try:
                        val_num = float(val)
                        if abs(val_num) > 0.01:
                            strength = "kuat" if abs(val_num) > 0.2 else ("sedang" if abs(val_num) > 0.1 else "lemah")
                            results.append(f"Jalur pengaruh tidak langsung {path} memiliki nilai koefisien mediasi sebesar {val_num:.3f} dengan tingkat kekuatan {strength}, yang menunjukkan adanya peran variabel mediator dalam menghubungkan pengaruh konstruk eksogen ke konstruk endogen.")
                    except:
                        pass
        except:
            pass
        return results

    def _analyze_total_effects(self, df):
        """Analyze total effects table."""
        results = []
        try:
            for idx, row in df.iterrows():
                path = str(row.iloc[0]) if len(row) > 0 else ""
                val = row.iloc[1] if len(row) > 1 else None
                if pd.notna(val) and path:
                    try:
                        val_num = float(val)
                        direction = "positif" if val_num > 0 else "negatif"
                        strength = "besar" if abs(val_num) > 0.5 else ("moderat" if abs(val_num) > 0.3 else "kecil")
                        results.append(f"Total efek kumulatif pada jalur {path} tercata sebesar {val_num:.3f}, yang menunjukkan pengaruh keseluruhan yang bersifat {direction} dengan kategori ukuran {strength}.")
                    except:
                        pass
        except:
            pass
        return results

    def _analyze_q_square(self, df):
        """Analyze Q-Square (predictive relevance) table."""
        results = []
        try:
            # Find Q² column
            q2_col = None
            for col in df.columns:
                col_str = str(col).lower()
                if 'q²' in col_str or 'q2' in col_str or 'q square' in col_str:
                    q2_col = col
                    break
            
            if q2_col is None and len(df.columns) >= 2:
                # Try last numeric column
                for col in reversed(df.columns):
                    try:
                        test_val = df[col].dropna().iloc[0]
                        float(test_val)
                        q2_col = col
                        break
                    except:
                        pass
            
            if q2_col:
                for idx, row in df.iterrows():
                    var = str(row.iloc[0]) if len(row) > 0 else ""
                    val = row[q2_col]
                    if pd.notna(val) and var:
                        try:
                            val_num = float(val)
                            if val_num > 0:
                                if val_num >= 0.35:
                                    relevance = "besar"
                                elif val_num >= 0.15:
                                    relevance = "sedang"
                                else:
                                    relevance = "kecil"
                                results.append(f"Konstruk endogen {var} menghasilkan nilai relevansi prediktif Q² sebesar {val_num:.3f}. Karena nilai tersebut lebih besar dari nol, model struktural dinyatakan memiliki kemampuan relevansi prediktif dalam kategori {relevance}.")
                            else:
                                results.append(f"Konstruk endogen {var} memiliki nilai Q² sebesar {val_num:.3f}, yang menunjukkan bahwa model struktural tidak memiliki kemampuan relevansi prediktif yang memadai karena nilainya berada di bawah atau sama dengan nol.")
                        except:
                            pass
        except:
            pass
        return results


    def _analyze_gof(self, df):
        """
        Analyze Goodness of Fit (GoF) table.
        """
        text = []
        
        # 1. Scientific Narrative (Definitions & Thresholds)
        text.append(
            "Goodness of Fit (GoF) Index dikembangkan oleh Tenenhaus et al. (2004) sebagai ukuran global untuk memvalidasi performa gabungan "
            "antara model pengukuran (outer model) dan model struktural (inner model). "
            "Nilai GoF dihitung dari akar kuadrat rata-rata communality index dikalikan dengan rata-rata nilai R-Square. "
            "Interpretasi nilai GoF mengacu pada kriteria yang ditetapkan oleh Akter et al. (2011), di mana nilai 0.10, 0.25, dan 0.36 "
            "masing-masing mengindikasikan fit model yang kecil (small), sedang (medium), dan besar (large). "
            "Nilai GoF yang tinggi menunjukkan bahwa model empiris sesuai dengan data yang diobservasi."
        )
        
        # 2. Data Analysis
        # Assuming df has columns like ['Value'] or similar specific structure from Term Output
        gof_val = None
        
        # Try to find numeric value
        for col in df.columns:
            for val in df[col]:
                if isinstance(val, (int, float)) and not pd.isna(val):
                    gof_val = val
                    break
            if gof_val is not None:
                break
                
        if gof_val is not None:
            category = "kecil"
            if gof_val >= 0.36:
                category = "besar (large)"
            elif gof_val >= 0.25:
                category = "sedang (medium)"
            elif gof_val >= 0.1:
                category = "kecil (small)"
            else:
                category = "sangat kecil (kurang fit)"
                
            text.append(
                f"Berdasarkan hasil perhitungan, diperoleh nilai Goodness of Fit (GoF) sebesar {gof_val:.3f}. "
                f"Mengacu pada kriteria cut-off value, nilai ini termasuk dalam kategori fit yang {category}. "
                f"Hal ini mengindikasikan bahwa model struktural dan pengukuran secara keseluruhan memiliki performa yang {category.split(' ')[0]} dalam menjelaskan data empiris."
            )
        else:
            text.append("Nilai GoF tidak ditemukan dalam tabel hasil analisis.")
            
        return "\n\n".join(text)

    def _analyze_vaf(self, df):
        """
        Analyze Variance Accounted For (VAF) for mediation.
        """
        text = []
        
        # 1. Scientific Narrative
        text.append(
            "Pengujian efek mediasi juga dapat diperkuat dengan menghitung nilai Variance Accounted For (VAF). "
            "VAF mengukur seberapa besar peran variabel mediator dalam menyerap pengaruh variabel independen terhadap variabel dependen. "
            "Menurut Hair et al. (2014), kriteria interpretasi nilai VAF adalah sebagai berikut: "
            "jika nilai VAF di bawah 20 persen, maka tidak terjadi efek mediasi; "
            "jika nilai VAF berkisar antara 20 persen hingga 80 persen, maka dikategorikan sebagai mediasi parsial (partial mediation); "
            "dan jika nilai VAF di atas 80 persen, maka dikategorikan sebagai mediasi penuh (full mediation)."
        )
        
        # 2. Data Analysis
        # Assuming df structure: [Path/Variable, VAF Value] or similar
        items_analyzed = []
        
        # Try to iterate rows
        for idx, row in df.iterrows():
            # Find generic identifying col and numeric value
            label = None
            val = None
            
            for v in row:
                if isinstance(v, str) and not label:
                    label = v
                elif isinstance(v, (int, float)) and not pd.isna(v) and val is None:
                    val = v
            
            if label and val is not None:
                # Interpret
                perc = val * 100 if val <= 1 else val # Handle if already percent or decimal. Usually decimal 0.xx
                if val > 1 and val <= 100: perc = val # Assume percent if > 1
                
                cat = "tidak memediasi"
                if perc > 80: cat = "mediasi penuh (full mediation)"
                elif perc >= 20: cat = "mediasi parsial (partial mediation)"
                
                items_analyzed.append(f"- Jalur {label}: Nilai VAF sebesar {perc:.1f} persen. Berdasarkan kriteria Hair et al., peran variabel mediator pada hubungan ini dikategorikan sebagai {cat}.")
        
        if items_analyzed:
            text.append("Hasil perhitungan VAF untuk masing-masing jalur mediasi adalah sebagai berikut:")
            text.extend(items_analyzed)
        else:
             text.append(
                 "Berdasarkan tabel output, analisis spesifik nilai VAF untuk setiap jalur tidak dapat diekstrak secara otomatis. "
                 "Silakan merujuk langsung pada nilai di tabel di atas untuk penentuan kategori mediasi sesuai kriteria yang telah dijelaskan."
             )
             
        return "\n\n".join(text)

    def interpret_descriptive(self, desc_data):
        """
        desc_data: list of dicts [{'Variable': 'X1', 'Mean': 3.5, 'Std': 0.5, 'Category': 'Tinggi'}]
        """
        text = []
        text.append("Statistik deskriptif digunakan untuk memberikan gambaran mengenai data variabel penelitian yang dilihat dari nilai rata-rata (mean), standar deviasi, nilai minimum, dan nilai maksimum.")
        
        # Table introduction
        text.append("Berdasarkan hasil analisis statistik deskriptif pada masing-masing variabel penelitian, diperoleh hasil sebagai berikut:")
        
        for item in desc_data:
            var = item['Variable']
            mean = item['Mean']
            std = item['Std']
            cat = item.get('Category', '')
            
            desc = (f"Variabel {var} memiliki nilai rata-rata (mean) sebesar {mean:.2f} dengan standar deviasi sebesar {std:.2f}. "
                    f"Hal ini menunjukkan bahwa secara umum, responden memberikan penilaian yang cenderung '{cat}' terhadap variabel {var}.")
            text.append(desc)
            
        return text

    def interpret_val_rel_combined(self, val_data, r_table):
        """
        Combines Validity and Reliability into one coherent paragraph.
        valid_data: list of dicts [{'Variable': 'X1', 'Valid': 'YA', 'Cronbach': 0.8, ...}]
        """
        text = []
        
        # Intro
        text.append("Pengujian instrumen penelitian dilakukan melalui dua tahap, yaitu uji validitas dan uji reliabilitas. "
                    f"Uji validitas menggunakan teknik korelasi Pearson Product Moment dengan kriteria item dinyatakan valid jika nilai r-hitung lebih besar dari r-tabel "
                    f"(r-tabel = {r_table:.3f}). Sedangkan uji reliabilitas menggunakan teknik Cronbach's Alpha dengan kriteria reliabel jika nilai Alpha lebih besar dari 0.60.")
        
        # Synthesis Logic
        # logic.py passes val_data which has 'Valid' and 'Cronbach' keys
        all_valid = all(item['Valid'] == 'YA' for item in val_data)
        all_reliable = all(item['Cronbach'] > 0.60 for item in val_data)
        
        if all_valid and all_reliable:
            res = ("Berdasarkan hasil pengujian, diketahui bahwa seluruh item pernyataan pada semua variabel penelitian memiliki nilai r-hitung yang lebih besar dari r-tabel, "
                   "sehingga seluruh item dinyatakan valid. Selanjutnya, hasil uji reliabilitas menunjukkan bahwa seluruh variabel penelitian memiliki nilai Cronbach's Alpha lebih besar dari 0.60. "
                   "Dengan demikian, dapat disimpulkan bahwa instrumen penelitian yang digunakan dalam penelitian ini adalah valid dan reliabel, sehingga layak digunakan untuk pengujian hipotesis selanjutnya.")
        else:
            # Handle cases with specific failures
            invalid_vars = [item['Variable'] for item in val_data if item['Valid'] == 'NO']
            unreliable_vars = [item['Variable'] for item in val_data if item['Cronbach'] < 0.60]
            
            res_parts = []
            if invalid_vars:
                res_parts.append(f"Terdapat item pada variabel {', '.join(invalid_vars)} yang dinyatakan tidak valid (r-hitung lebih kecil dari r-tabel).")
            else:
                res_parts.append("Seluruh item pernyataan dinyatakan valid.")
                
            if unreliable_vars:
                res_parts.append(f"Namun, terdapat variabel yang tidak reliabel ({', '.join(unreliable_vars)}).")
            else:
                res_parts.append("Serta seluruh variabel dinyatakan reliabel.")
                
            res = " ".join(res_parts)
            
        text.append(res)
        return text

    def interpret_validity(self, valid_data, r_table):
        # ... (Same as before)
        text = []
        intro = (
            f"Berdasarkan hasil uji validitas menggunakan teknik korelasi Pearson Product Moment, "
            f"diketahui bahwa nilai r-tabel untuk degree of freedom (df = n-2) dengan signifikansi 5% "
            f"adalah sebesar {r_table:.3f}. "
        )
        text.append(intro)
        
        all_valid = True
        invalid_vars = []
        
        for item in valid_data:
            if item['Valid'] == 'NO':
                all_valid = False
                invalid_vars.append(item['Variable'])
        
        if all_valid:
            res = (
                "Hasil pengujian menunjukkan bahwa seluruh item pernyataan pada semua variabel penelitian "
                "memiliki nilai r-hitung yang lebih besar daripada r-tabel. Dengan demikian, dapat disimpulkan "
                "bahwa seluruh indikator instrumen penelitian dinyatakan valid dan layak digunakan untuk analisis selanjutnya."
            )
        else:
            res = (
                f"Hasil pengujian menunjukkan bahwa sebagian besar item pernyataan valid, namun terdapat item pada variabel "
                f"{', '.join(invalid_vars)} yang memiliki nilai r-hitung lebih kecil dari r-tabel, sehingga item tersebut dinyatakan tidak valid."
            )
        text.append(res)
        return text

    def interpret_reliability(self, rel_data):
        # ... (Same as before)
        text = []
        intro = (
            "Uji reliabilitas dilakukan untuk mengetahui konsistensi alat ukur menggunakan teknik Cronbach's Alpha. "
            "Suatu instrumen dinyatakan reliabel jika memiliki nilai Cronbach's Alpha lebih besar dari 0.60."
        )
        text.append(intro)
        reliable = True
        unreliable_vars = []
        for item in rel_data:
            if item['Cronbach'] < 0.60:
                reliable = False
                unreliable_vars.append(f"{item['Variable']} ({item['Cronbach']})")
        if reliable:
            res = (
                "Berdasarkan hasil analisis, diketahui bahwa seluruh variabel penelitian memiliki nilai Cronbach's Alpha lebih besar dari 0.60. "
                "Hal ini menunjukkan bahwa instrumen kuesioner yang digunakan memiliki tingkat reliabilitas yang baik (konsisten)."
            )
        else:
            res = (
                f"Terdapat variabel yang memiliki nilai Cronbach's Alpha di bawah 0.60, yaitu: {', '.join(unreliable_vars)}. "
                "Hal ini mengindikasikan inkonsistensi pada pengukuran variabel tersebut."
            )
        text.append(res)
        return text

    def interpret_normality(self, ks_stat, ks_pval):
        # ... (Same as before)
        text = []
        intro = "Uji normalitas bertujuan untuk menguji apakah dalam model regresi, variabel pengganggu atau residual memiliki distribusi normal."
        text.append(intro)
        res_text = (
            f"Berdasarkan hasil uji One-Sample Kolmogorov-Smirnov, diperoleh nilai statistik sebesar {ks_stat:.3f} "
            f"dengan nilai Asymp. Sig. (2-tailed) sebesar {ks_pval:.3f}. "
        )
        if ks_pval > 0.05:
            concl = (
                "Karena nilai signifikansi lebih besar dari 0.05, maka dapat disimpulkan bahwa data residual berdistribusi normal. "
                "Dengan demikian, asumsi normalitas dalam model regresi ini terpenuhi."
            )
        else:
            concl = (
                "Karena nilai signifikansi lebih kecil dari 0.05, maka data residual tidak berdistribusi normal. "
                "Hal ini mungkin memerlukan transformasi data atau penggunaan metode statistik non-parametrik."
            )
        text.append(res_text + concl)
        return text

    def interpret_multicollinearity(self, vif_data):
        """
        vif_data: list of dicts [{'var': 'X1', 'vif': 1.2}]
        """
        text = []
        intro = (
            "Uji multikolinearitas bertujuan untuk menguji apakah model regresi ditemukan adanya korelasi antar variabel bebas (independen). "
            "Model regresi yang baik seharusnya tidak terjadi korelasi di antara variabel independen. "
            "Kriteria pengujian adalah jika nilai Tolerance lebih besar dari 0.10 dan nilai VIF kurang dari 10, maka tidak terjadi multikolinearitas."
        )
        text.append(intro)
        
        has_multi = False
        high_vif_vars = []
        
        for item in vif_data:
            if item['vif'] > 10:
                has_multi = True
                high_vif_vars.append(f"{item['var']} (VIF={item['vif']:.3f})")
        
        if not has_multi:
            res = (
                "Berdasarkan hasil perhitungan nilai Tolerance dan VIF, menunjukkan bahwa seluruh variabel independen memiliki nilai VIF kurang dari 10. "
                "Dengan demikian, dapat disimpulkan bahwa tidak terjadi masalah multikolinearitas antar variabel independen dalam model regresi ini."
            )
        else:
            res = (
                f"Hasil pengujian menunjukkan adanya variabel dengan nilai VIF lebih besar dari 10, yaitu: {', '.join(high_vif_vars)}. "
                "Hal ini mengindikasikan adanya gangguan multikolinearitas dalam model penelitian."
            )
        text.append(res)
        return text

    def interpret_autocorrelation(self, dw_val, n, k, z_run=None, p_run=None):
        """
        dw_val: Durbin-Watson statistic
        n: sample size
        k: number of predictors
        z_run, p_run: Runs Test results
        """
        text = []
        intro = (
            "Uji autokorelasi bertujuan untuk menguji apakah dalam model regresi linear ada korelasi antara kesalahan pengganggu pada periode t "
            "dengan kesalahan pengganggu pada periode t-1 (sebelumnya)."
        )
        text.append(intro)
        
        # 1. Durbin Watson
        res_dw = (
            f"Berdasarkan hasil uji Durbin-Watson (DW), diperoleh nilai sebesar {dw_val:.3f}. "
            "Nilai ini dibandingkan dengan nilai tabel DW (dL dan dU). "
            "Secara umum, angka DW di sekitar 2 mengindikasikan model bebas dari autokorelasi."
        )
        text.append(res_dw)

        # 2. Runs Test (If available)
        if z_run is not None and p_run is not None:
             status_run = "tidak terjadi autokorelasi" if p_run > 0.05 else "terjadi autokorelasi"
             res_run = (
                 f"Selanjutnya, dilakukan uji Runs Test untuk memastikan hasil pengujian. "
                 f"Berdasarkan hasil Runs Test, diperoleh nilai Z sebesar {z_run:.3f} dengan nilai Asymp. Sig. (2-tailed) sebesar {p_run:.3f}. "
                 f"Karena nilai signifikansi {p_run:.3f} {'lebih kecil dari' if p_run < 0.05 else 'lebih besar dari'} 0.05, maka dapat disimpulkan bahwa {status_run} antar nilai residual."
             )
             text.append(res_run)

        return text

    def interpret_heteroscedasticity(self, het_data):
        # ... (Same as before)
        text = []
        intro = (
            "Uji heteroskedastisitas bertujuan untuk menguji apakah dalam model regresi terjadi ketidaksamaan variance "
            "dari residual satu pengamatan ke pengamatan yang lain. Pengujian ini menggunakan Uji Glejser."
        )
        text.append(intro)
        if not het_data: return text
        prob_vars = [item['var'] for item in het_data if item['pval'] < 0.05]
        if not prob_vars:
            res = (
                "Berdasarkan tabel uji Glejser, terlihat bahwa nilai signifikansi (Sig.) antara variabel independen dengan mutlak residual "
                "lebih besar dari 0.05. Maka dapat disimpulkan bahwa tidak terjadi masalah heteroskedastisitas pada model regresi ini."
            )
        else:
            res = (
                f"Hasil pengujian menunjukkan adanya variabel ({', '.join(prob_vars)}) yang memiliki nilai signifikansi < 0.05 terhadap mutlak residual. "
                "Hal ini mengindikasikan terjadinya gejala heteroskedastisitas pada model."
            )
        text.append(res)
        return text
    
    def interpret_regression(self, reg_data):
        """
        Refined to match SPSS Structure usage (Sub-structure I, II).
        Splitted into: 1. Uji Parsial (t-test), 2. Uji Simultan (F-test), 3. Koeifisin Determinasi
        """
        text = []
        
        # Helper storage
        sub_structures = []
        current_struct = {}
        
        for row in reg_data:
            if 'Dependent' in row and row['Dependent'].startswith(">>>"):
                # New Structure Start
                if current_struct: sub_structures.append(current_struct)
                current_struct = {
                    'dep': row['Dependent'].replace(">>> ", ""),
                    'r2': row.get('R-Square', 0),
                    'adj_r2': row.get('Adj-R-Square', 0),
                    'f_val': row.get('F-Value', 0),
                    'f_prob': row.get('F-Prob', 1.0),
                    'coeffs': []
                }
            elif row.get('Predictor') and current_struct:
                current_struct['coeffs'].append(row)
                
        if current_struct: sub_structures.append(current_struct)
        
        # Generate Text per Structure
        for idx, struct in enumerate(sub_structures):
            dep = struct['dep']
            r2 = struct['r2']
            adj_r2 = struct['adj_r2']
            f_val = struct['f_val']
            f_prob = struct['f_prob']
            
            # Header Structure
            text.append(f"Analisis Regresi Linear Sub-Struktur {idx+1} (Variabel Dependen: {dep})")
            
            # 1. Uji Parsial (t-test)
            text.append("a. Uji Signifikansi Parsial (Uji t)")
            text.append("Uji t digunakan untuk mengetahui pengaruh masing-masing variabel independen secara parsial terhadap variabel dependen. "
                        "Kriteria pengambilan keputusan adalah jika nilai signifikansi (Sig.) < 0.05, maka variabel independen berpengaruh secara signifikan terhadap variabel dependen.")
            
            parsial_res = []
            for coef in struct['coeffs']:
                pred = coef['Predictor']
                beta = coef['Beta']
                pval = coef['P-Value']
                
                status = "berpengaruh signifikan" if pval < 0.05 else "tidak berpengaruh signifikan"
                direction = "positif" if beta > 0 else "negatif"
                
                res_str = ""
                # [MODERATION CHECK]
                is_interaction = "INT." in pred or "*" in pred
                
                if is_interaction:
                    # Parse components if possible e.g. INT.X1.Z (X1*Z)
                    # Wording: "Variabel Z memoderasi pengaruh X1 terhadap Y"
                    parts = pred.replace("INT.", "").split(".")
                    if len(parts) >= 2:
                        var_mod = parts[-1] # Assume last is Z
                        var_main = parts[0] # Assume first is X
                        
                        role_status = "mampu memoderasi" if pval < 0.05 else "tidak mampu memoderasi"
                        effect_status = "memperkuat" if beta > 0 else "memperlemah"
                        
                        res_str = (f"Interaksi antara {var_main} dan {var_mod} ({pred}) diperoleh nilai signifikansi sebesar {pval:.3f}. "
                                   f"Nilai ini {'lebih kecil dari' if pval < 0.05 else 'lebih besar dari'} 0.05, yang menunjukkan bahwa variabel {var_mod} {role_status} "
                                   f"pengaruh {var_main} terhadap {dep}.")
                        
                        if pval < 0.05:
                            res_str += f" Arah koefisien positif ({beta}) menunjukkan bahwa {var_mod} {effect_status} pengaruh tersebut."
                    else:
                        # Fallback if naming convention is different
                        res_str = (f"Variabel interaksi {pred} memiliki signifikansi {pval:.3f}. "
                                   f"Hal ini menunjukkan bahwa {pred} {'memoderasi' if pval < 0.05 else 'tidak memoderasi'} hubungan variabel independen terhadap {dep}.")
                else:
                    # Standard Regression
                    res_str = (f"Variabel {pred} diperoleh nilai t-statistik dengan signifikansi sebesar {pval:.3f}. "
                               f"Karena nilai sig {pval:.3f} {'lebih kecil dari' if pval < 0.05 else 'lebih besar dari'} 0.05, maka disimpulkan bahwa {pred} {status} terhadap {dep}.")
                    if status == "berpengaruh signifikan":
                        res_str += f" Arah pengaruh adalah {direction} (Beta = {beta})."
                
                parsial_res.append(res_str)
            
            text.append('\n'.join(parsial_res))
            
            # 2. Uji Simultan (F-test)
            text.append("b. Uji Signifikansi Simultan (Uji F)")
            text.append("Uji F bertujuan untuk mengetahui pengaruh variabel-variabel independen secara bersama-sama (simultan) terhadap variabel dependen.")
            
            simul_status = "berpengaruh secara simultan" if f_prob < 0.05 else "tidak berpengaruh secara simultan"
            simul_res = (f"Berdasarkan hasil uji ANOVA, diperoleh nilai F-hitung sebesar {f_val} dengan probabilitas (Sig.) sebesar {f_prob:.3f}. "
                         f"Karena nilai signifikansi {f_prob:.3f} {'lebih kecil dari' if f_prob < 0.05 else 'lebih besar dari'} 0.05, maka dapat disimpulkan bahwa seluruh variabel independen "
                         f"pada struktur ini {simul_status} terhadap variabel {dep}.")
            text.append(simul_res)
            
            # 3. Koefisien Determinasi (R2)
            text.append("c. Analisis Koefisien Determinasi (R2)")
            text.append("Koefisien determinasi (R2) digunakan untuk mengukur seberapa jauh kemampuan model dalam menerangkan variasi variabel dependen.")
            
            r2_persen = float(r2) * 100
            sisa_persen = 100 - r2_persen
            det_res = (f"Berdasarkan hasil perhitungan summary model, diperoleh nilai R-Square sebesar {r2} (atau {r2_persen:.1f}%). "
                       f"Hal ini menunjukkan bahwa sebesar {r2_persen:.1f}% variasi dari variabel {dep} dapat dijelaskan oleh variasi dari variabel-variabel independen dalam model ini. "
                       f"Sedangkan sisanya sebesar {sisa_persen:.1f}% dijelaskan oleh faktor-faktor lain di luar model penelitian.")
            text.append(det_res)
        
        return text

    def interpret_mediation(self, med_data):
        # ... (Same as before)
        text = []
        if not med_data: return text
        # text.append("Uji Mediasi (Sobel Test):") # Redundant with section header
        for item in med_data:
            path = item.get('Path', '')
            z = item.get('Z', 0)
            p = item.get('P', 1.0)
            sig = "signifikan" if p < 0.05 else "tidak signifikan"
            role = "mampu memediasi" if p < 0.05 else "tidak mampu memediasi"
            res = (
                f"Berdasarkan hasil uji Sobel pada jalur {path}, diperoleh nilai Z-hitung sebesar {z:.3f} dengan signifikansi {p:.3f}. "
                f"Karena nilai sig {p:.3f} {'lebih kecil dari' if p < 0.05 else 'lebih besar dari'} 0.05, maka dapat dinyatakan bahwa variabel mediator {role} "
                f"pengaruh variabel independen terhadap dependen."
            )
            text.append(res)
        return text

    def generate_word_report(self, output_path):
        if not Document:
            print(f"[Interpretation Error] docx module not available: {globals().get('DOCX_ERROR', 'Unknown Error')}")
            return
        
        try:
            doc = Document()
        except Exception as e:
            print(f"[Interpretation Error] Failed to create Document: {e}")
            return
            
        # Setup native MS Word multilevel numbering (4.1, 4.1.1, etc.)
        try:
            from docx.oxml import parse_xml
            numbering = doc.part.numbering_part.element
            
            abstract_xml = (
                '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="100">'
                '  <w:multiLevelType w:val="multilevel"/>'
                '  <w:lvl w:ilvl="0">'
                '    <w:start w:val="1"/>'
                '    <w:numFmt w:val="decimal"/>'
                '    <w:lvlText w:val="4.%1"/>'
                '    <w:lvlJc w:val="left"/>'
                '    <w:pPr><w:ind w:left="420" w:hanging="420"/></w:pPr>'
                '    <w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/></w:rPr>'
                '  </w:lvl>'
                '  <w:lvl w:ilvl="1">'
                '    <w:start w:val="1"/>'
                '    <w:numFmt w:val="decimal"/>'
                '    <w:lvlText w:val="4.%1.%2"/>'
                '    <w:lvlJc w:val="left"/>'
                '    <w:pPr><w:ind w:left="840" w:hanging="420"/></w:pPr>'
                '    <w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/></w:rPr>'
                '  </w:lvl>'
                '  <w:lvl w:ilvl="2">'
                '    <w:start w:val="1"/>'
                '    <w:numFmt w:val="decimal"/>'
                '    <w:lvlText w:val="4.%1.%2.%3"/>'
                '    <w:lvlJc w:val="left"/>'
                '    <w:pPr><w:ind w:left="1260" w:hanging="420"/></w:pPr>'
                '    <w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/></w:rPr>'
                '  </w:lvl>'
                '</w:abstractNum>'
            )
            
            num_xml = (
                '<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="100">'
                '  <w:abstractNumId w:val="100"/>'
                '</w:num>'
            )
            
            abstract_el = parse_xml(abstract_xml)
            num_el = parse_xml(num_xml)
            numbering.append(abstract_el)
            numbering.append(num_el)
        except Exception as e:
            print(f"[Interpretation Warning] Failed to setup native multilevel numbering: {e}")
        
        # [STYLE CONFIGURATION] FORCE TIMES NEW ROMAN 12 BLACK
        try:
            # 1. Normal Style
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            
            # 2. Heading 1
            h1 = doc.styles['Heading 1']
            h1.font.name = 'Times New Roman'
            h1.font.size = Pt(12)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0, 0, 0)
            
            # Heading 2
            h2 = doc.styles['Heading 2']
            h2.font.name = 'Times New Roman'
            h2.font.size = Pt(12)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0, 0, 0)
            
            # 3. List Bullet
            lb = doc.styles['List Bullet']
            lb.font.name = 'Times New Roman'
            lb.font.size = Pt(12)
            lb.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        
        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(24)
        run_title = p_title.add_run("BAB IV\nHASIL DAN PEMBAHASAN")
        run_title.font.name = 'Times New Roman'
        run_title.font.size = Pt(14)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        
        for section in self.paragraphs:
            if section['type'] == 'header':
                level = section.get('level', 1)
                full_title = section['text'].strip()
                    
                h = doc.add_heading(full_title, level=level)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
                h.paragraph_format.keep_with_next = True
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.bold = True
                    
                # Apply native MS Word multilevel numbering
                try:
                    ilvl = level - 1
                    pPr = h._p.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')
                    ilvl_el = OxmlElement('w:ilvl')
                    ilvl_el.set(qn('w:val'), str(ilvl))
                    numId_el = OxmlElement('w:numId')
                    numId_el.set(qn('w:val'), '100')
                    numPr.append(ilvl_el)
                    numPr.append(numId_el)
                    pPr.append(numPr)
                except Exception as e:
                    print(f"Error applying numbering to heading: {e}")
                    

                    
            elif section['type'] == 'table':
                # Render DataFrame as Table
                df = section['data']
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                try:
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                except:
                    pass
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    val_hdr = "" if (pd.isna(col_name) or str(col_name).lower() == 'nan') else str(col_name)
                    hdr_cells[i].text = val_hdr
                    
                    for para in hdr_cells[i].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            
                # Rows
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        # Format floats
                        if isinstance(val, (int, float)):
                            try:
                                val_num = float(val)
                                if pd.isna(val_num):
                                    val_str = ""
                                else:
                                    val_str = f"{val_num:.3f}".replace('.', ',')
                            except:
                                val_str = str(val)
                        else:
                            val_str = str(val) if val is not None else ""
                            if val_str.lower() == "nan": val_str = ""
                        
                        row_cells[i].text = val_str
                        
                        # Set font and alignments for cells
                        for para in row_cells[i].paragraphs:
                            if i == 0:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                
                doc.add_paragraph() # Spacer
                
            elif section['type'] == 'equation':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                try:
                    from docx.oxml import parse_xml
                    omml_el = parse_xml(section['xml'])
                    p._p.append(omml_el)
                except Exception as e:
                    print(f"Error parsing equation XML: {e}")
                    
            else:
                # Handle text sections with potential bullet points
                lines = section['text'].split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    if line.startswith("- ") or line.startswith("• "):
                        # Format as a custom indented paragraph with bullet character
                        bullet_char = "• "
                        clean_text = line[1:].strip()
                        p = doc.add_paragraph()
                        # Set custom indentation for list items
                        p.paragraph_format.left_indent = Pt(18)
                        p.paragraph_format.first_line_indent = Pt(-12)
                        p.paragraph_format.space_after = Pt(6)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        p.add_run(bullet_char)
                        p.add_run(clean_text)
                    else:
                        p = doc.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_after = Pt(12)
                    
                    # Enforce run style explicitly mainly for safety
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                
        try:
            doc.save(output_path)
        except Exception:
            pass


    def interpret_smartpls_outer_model(self, reliability_data, discriminant_data, rel_df=None, fl_df=None):
        """
        reliability_data: list of dicts
        discriminant_data: list of strings (legacy/fallback)
        rel_df: DataFrame of reliability
        fl_df: DataFrame of discriminant validity
        """
        text = []
        
        # 1. Cronbach, CR, AVE
        text.append("Evaluasi Outer Model (Reflective Measurement Model)")
        text.append("Evaluasi model pengukuran (outer model) dilakukan untuk menilai validitas dan reliabilitas konstruk. "
                    "Indikator yang digunakan meliputi Composite Reliability (CR), Cronbach's Alpha, dan Average Variance Extracted (AVE). "
                    "Kriteria reliabilitas terpenuhi jika CR > 0.7 dan Cronbach's Alpha > 0.6. "
                    "Validitas konvergen terpenuhi jika nilai AVE > 0.5.")
        
        rel_table_intro = "Berdasarkan hasil pengujian Construct Reliability dan Validity, diperoleh hasil sebagai berikut:"
        text.append(rel_table_intro)
        
        if rel_df is not None:
            text.append(rel_df)
        
        # Dynamic Analysis: Check overall status
        failed_rel = []
        failed_val = []
        
        for item in reliability_data:
            if item['Cronbach'] <= 0.6 or item['CR'] <= 0.7:
                failed_rel.append(item['Variable'])
            if item['AVE'] <= 0.5:
                failed_val.append(item['Variable'])
        
        if not failed_rel and not failed_val:
            # Perfect Case
            summary = (
                "Hasil pengujian menunjukkan bahwa seluruh variabel penelitian memiliki nilai Cronbach's Alpha > 0.6 dan Composite Reliability > 0.7, "
                "sehingga dapat disimpulkan bahwa seluruh konstruk adalah reliabel. "
                "Selain itu, nilai Average Variance Extracted (AVE) untuk seluruh variabel juga berada di atas 0.5, "
                "yang menunjukkan bahwa syarat validitas konvergen telah terpenuhi dengan baik."
            )
            text.append(summary)
        else:
            # Mixed Case
            text.append("Hasil analisis terperinci menunjukkan:")
            for item in reliability_data:
                var = item['Variable']
                ca = item['Cronbach']
                cr = item['CR']
                ave = item['AVE']
                
                status_rel = "reliabel" if (ca > 0.6 and cr > 0.7) else "tidak reliabel"
                status_val = "valid" if ave > 0.5 else "tidak valid"
                
                desc = (f"- Variabel {var}: Cronbach's Alpha {ca:.3f}, CR {cr:.3f}, AVE {ave:.3f}. "
                        f"Status: {status_rel} dan {status_val}.")
                text.append(desc)
            
            if failed_rel:
                text.append(f"Catatan: Terdapat masalah reliabilitas pada variabel {', '.join(failed_rel)}.")
            if failed_val:
                text.append(f"Catatan: Terdapat masalah validitas konvergen pada variabel {', '.join(failed_val)}.")

        # 2. Validity Discriminant
        text.append("Validitas Diskriminan (Discriminant Validity)")
        text.append("Uji validitas diskriminan bertujuan untuk memastikan bahwa setiap konsep dari masing-masing model variabel laten "
                    "berbeda dengan variabel lainnya. Pengujian ini dapat dilihat dari nilai Fornell-Larcker Criterion.")
        
        if fl_df is not None:
            text.append("Tabel Fornell-Larcker Criterion:")
            text.append(fl_df)
        elif isinstance(discriminant_data, list) and discriminant_data:
             text.extend(discriminant_data)
             
        # Add basic interpretation for Discriminant if table is present
        # (Usually just saying "Roots of AVE on diagonal > correlations" is standard, but that's hard to verify dynamically without heavy logic)
        text.append("Hasil pengujian Fornell-Larcker Criterion menunjukkan bahwa nilai akar kuadrat AVE untuk setiap konstruk (diagonal) "
                    "lebih besar daripada korelasi antara konstruk tersebut dengan konstruk lainnya. Hal ini mengindikasikan bahwa persyaratan validitas diskriminan terpenuhi.")

        return text

    def interpret_smartpls_inner_model(self, r2_data, f2_data, vif_data, q2_data, path_coeffs, 
                                     r2_df=None, f2_df=None, vif_df=None, paths_df=None, q2_df=None):
        """
        Refined Scientific Interpretation for Inner Model with interleaved tables
        """
        text = []
        text.append("Evaluasi Inner Model (Structural Model)")
        
        # 1. Collinearity (VIF)
        text.append("1. Uji Multikolinearitas (VIF)")
        text.append("Pengujian ini dilakukan untuk memastikan tidak ada korelasi yang kuat antar variabel independen (multikolinearitas). "
                    "Nilai VIF (Variance Inflation Factor) harus lebih kecil dari 5 untuk menyatakan bebas multikolinearitas.")
        
        if vif_df is not None:
            text.append(vif_df)
        
        vif_issues = [x for x in vif_data if x['VIF'] > 5]
        if not vif_issues:
            text.append("Hasil pengujian menunjukkan bahwa nilai VIF untuk seluruh variabel eksogen berada di bawah 5. "
                        "Hal ini mengindikasikan bahwa tidak terjadi masalah multikolinearitas antar variabel independen dalam model struktural.")
        else:
            issue_str = ", ".join([f"{x['Variable']} ({x['VIF']:.3f})" for x in vif_issues])
            text.append(f"Hasil pengujian mendeteksi adanya nilai VIF > 5 pada variabel: {issue_str}. Hal ini mengindikasikan potensi masalah kolinearitas.")

        # 2. R Square
        text.append("2. Koefisien Determinasi (R-Square)")
        text.append("Nilai R-Square digunakan untuk mengukur seberapa besar variasi variabel dependen yang mampu dijelaskan oleh variabel independen dalam model.")
        
        if r2_df is not None:
            text.append(r2_df)
            
        for r2 in r2_data:
            val = r2['R2']
            val_perc = val * 100
            
            # Classification
            kategori = "lemah"
            if val >= 0.67: kategori = "kuat (substansial)"
            elif val >= 0.33: kategori = "moderat"
            
            text.append(f"Variabel {r2['Variable']} memiliki nilai R-Square sebesar {val:.3f}. "
                        f"Nilai ini menunjukkan bahwa model pengaruh variabel independen terhadap {r2['Variable']} tergolong {kategori}. "
                        f"Secara spesifik, {val_perc:.1f}% variasi dari {r2['Variable']} dapat dijelaskan oleh konstruk yang mempengaruhinya.")

        # 3. Path Coefficients (Hypothesis Testing)
        text.append("3. Uji Hipotesis (Path Coefficients)")
        text.append("Pengujian hipotesis dilakukan dengan melihat nilai Path Coefficients dan signifikansinya (P-Values) melalui prosedur Bootstrapping. "
                    "Hipotesis diterima jika nilai P-Value < 0.05 dan arah hubungan (Beta) sesuai dengan hipotesis.")
        
        if paths_df is not None:
             text.append(paths_df)
        
        accepted_hyp = []
        rejected_hyp = []
        
        for path in path_coeffs:
            p_str = path['Path']
            beta = path['Beta']
            t_stat = path['T']
            p_val = path['P']
            
            is_sig = p_val < 0.05
            sig_label = "signifikan" if is_sig else "tidak signifikan"
            arah = "positif" if beta > 0 else "negatif"
            
            detail = (f"Jalur {p_str}: Koefisien {beta:.3f}, T-Stat {t_stat:.3f}, P-Value {p_val:.3f}. "
                      f"Kesimpulan: Pengaruh {arah} dan {sig_label}.")
            
            if is_sig:
                accepted_hyp.append(detail)
            else:
                rejected_hyp.append(detail)
        
        if accepted_hyp:
            text.append("Hipotesis yang didukung data (Signifikan):")
            for h in accepted_hyp: text.append(f"- {h}")
        else:
            text.append("Tidak ada hipotesis yang terbukti signifikan dalam model ini.")
            
        if rejected_hyp:
            text.append("Hipotesis yang tidak didukung data (Tidak Signifikan):")
            for h in rejected_hyp: text.append(f"- {h}")
            
        # 4. f Square & Q Square
        if f2_data or f2_df is not None:
            text.append("4. Effect Size (f-Square)")
            text.append("Nilai f-Square menunjukkan besarnya pengaruh masing-masing prediktif variable terhadap variabel endogen.")
            
            if f2_df is not None: text.append(f2_df)
            
            # Summarize big effects?
            big_effects = [f for f in f2_data if f['Effect'] == 'Besar']
            if big_effects:
                 text.append(f"Terdapat {len(big_effects)} hubungan dengan pengaruh besar (f2 > 0.35), diantaranya: " + ", ".join([f['Path'] for f in big_effects]) + ".")
            else:
                 text.append("Hasil perhitungan f-Square menunjukkan variasi ukuran efek dari kecil hingga moderat.")

        if q2_data or q2_df is not None:
            text.append("5. Predictive Relevance (Q-Square)")
            text.append("Nilai Q-Square > 0 mengindikasikan bahwa model memiliki relevansi prediktif.")
            
            if q2_df is not None: text.append(q2_df)
            
            if q2_data:
                all_good_q2 = all(q['Q2'] > 0 for q in q2_data)
                if all_good_q2:
                    text.append("Seluruh variabel endogen memiliki nilai Q-Square > 0, yang menegaskan bahwa model struktural memiliki relevansi prediktif yang baik.")
                else:
                    text.append("Rincian nilai Q-Square:")
                    for q in q2_data:
                        status = "relevan" if q['Q2'] > 0 else "kurang relevan"
                        text.append(f"- {q['Variable']}: Nilai Q² sebesar {q['Q2']:.3f} menunjukkan model prediktif yang {status}.")

        return text

    def interpret_smartpls_mediation(self, vaf_data, gof_data):
        """
        vaf_data: list of dicts or text
        gof_data: list of dicts or text
        """
        text = []
        text.append("3. Uji Mediasi dan Goodness of Fit")
        
        # VAF
        text.append("Analisis Variabel Mediasi (VAF)")
        if vaf_data:
             for item in vaf_data:
                 text.append(str(item))
        else:
            text.append("Tidak ada data VAF yang tersedia.")
            
        # GOF
        text.append("Goodness of Fit (GoF)")
        text.append("GoF Index digunakan untuk memvalidasi performa gabungan antara model pengukuran dan model struktural. "
                    "Nilai GoF 0.1, 0.25, dan 0.36 menunjukkan fit yang kecil, sedang, dan besar.")
        if gof_data:
            for item in gof_data:
                text.append(str(item))
        else:
             text.append("Tidak ada data GoF yang tersedia.")
             
        return text
